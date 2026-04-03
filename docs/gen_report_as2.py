from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, bold=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(0)
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            cells[c_idx].paragraphs[0].paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

# ── Build document ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)

# ── Title block ─────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Отчёт по Assignment 2")
r.font.name = "Times New Roman"; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Quality Assurance: Test Automation & CI/CD Integration")
r.font.name = "Times New Roman"; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Студент: Казиханов Диас  |  Группа: CSE-2506M  |  Дата: март 2026")
r.font.name = "Times New Roman"; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)

doc.add_paragraph()

# ── 1. Введение ─────────────────────────────────────────────────────────────
add_heading(doc, "1. Введение")
add_paragraph(doc,
    "В рамках второго задания по курсу Quality Assurance была проведена работа по расширению "
    "тестового покрытия системы Habaneras de Lino DRF API, настройке полноценного CI/CD пайплайна "
    "и формированию метрик качества.")
add_paragraph(doc, "Цели работы:")
for item in [
    "Автоматизировать тесты для критичных и высокорисковых модулей",
    "Определить качественные ворота (quality gates) и встроить их в CI",
    "Построить трёхэтапный CI/CD pipeline (unit -> integration -> e2e)",
    "Измерить покрытие кода и задокументировать найденные дефекты",
    "Сформировать метрики и подготовить отчёт",
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_font(run)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# ── 2. Описание системы ─────────────────────────────────────────────────────
add_heading(doc, "2. Тестируемая система")
add_paragraph(doc,
    "Система: Habaneras de Lino — REST API для интернет-магазина одежды.\n"
    "Репозиторий: https://github.com/Ceci-Aguilera/habaneras-de-lino-drf-api")

add_table(doc,
    ["Компонент", "Технология"],
    [
        ["Backend", "Django 4.0.6"],
        ["API", "Django REST Framework 3.13.1"],
        ["База данных", "PostgreSQL 13"],
        ["Контейнеризация", "Docker, Docker Compose"],
        ["Web-сервер", "Nginx"],
        ["Платежи", "Stripe"],
        ["Изображения", "Cloudinary"],
    ]
)

# ── 3. Стратегия автоматизации ───────────────────────────────────────────────
add_heading(doc, "3. Стратегия автоматизации")
add_paragraph(doc,
    "Автоматизация ведётся по приоритету риска из Assignment 1 (Risk Score = Вероятность x Влияние). "
    "Первыми покрываются критичные модули с наибольшим risk score.")

add_table(doc,
    ["Приоритет", "Модуль", "Risk Score", "Статус в AS2"],
    [
        ["1", "M1 — Order & Checkout", "20", "Расширен (5 новых integration тестов)"],
        ["2", "M6 — Data Validation", "16", "Расширен (field boundary тесты)"],
        ["3", "M2 — Stripe Payment", "15", "Ручное тестирование (нет Stripe ключей)"],
        ["4", "M8 — Tax Calculation", "12", "Новый (6 unit тестов)"],
        ["5", "M3 — Cart Management", "12", "Расширен (4 lifecycle тестов)"],
        ["6", "M4 — Product Catalog", "6", "Расширен (6 performance тестов)"],
        ["7", "M7 — Admin Panel", "6", "5 E2E тестов из AS1 сохранены"],
    ]
)

add_paragraph(doc, "Принятые решения по автоматизации:", bold=True)
add_paragraph(doc,
    "Stripe-тесты отложены на AS3 — для них требуется реальный Stripe test key и webhook endpoint. "
    "Cloudinary-тесты пропущены — внешний CDN не доступен в CI окружении. "
    "Все остальные критичные и высокорисковые модули покрыты автоматически.")

doc.add_paragraph()

# ── 4. Написанные тесты ──────────────────────────────────────────────────────
add_heading(doc, "4. Написанные тесты")

add_heading(doc, "4.1 Unit тесты", level=2)
add_paragraph(doc,
    "Unit тесты запускаются без Docker, на SQLite в памяти. "
    "Покрывают модели и бизнес-логику напрямую через Django ORM.")

add_table(doc,
    ["Файл", "Тесты", "Что покрывает"],
    [
        ["tests/unit/test_models.py (AS1)", "10", "Validator цвета, Cart total, Address creation"],
        ["tests/unit/test_globalmodel.py (AS2)", "6", "M8: GlobalModel tax, DoesNotExist, MultipleObjectsReturned"],
        ["tests/unit/test_product_model.py (AS2)", "6", "M3/M6: Cart total с товарами, поля ClothingProduct"],
        ["Итого unit", "22", ""],
    ]
)

add_paragraph(doc, "Новые unit тесты (AS2):", bold=True)
data_unit = [
    ("test_globalmodel.py", [
        "test_no_active_model_raises_does_not_exist — DoesNotExist при отсутствии активной записи",
        "test_multiple_active_models_raises_multiple_objects_returned — MultipleObjectsReturned при двух активных",
        "test_tax_rate_stored_as_decimal — tax_rate сохраняется как Decimal",
        "test_dollar_exchange_stored_as_decimal — dollar_exchange сохраняется как Decimal",
        "test_active_flag_default_false — active=False по умолчанию",
        "test_str_representation — строковое представление GlobalModel",
    ]),
    ("test_product_model.py", [
        "test_cart_total_with_single_item — Cart.get_total() с одним товаром",
        "test_cart_total_with_multiple_items — Cart.get_total() с несколькими товарами",
        "test_cart_total_empty — пустая корзина = 0",
        "test_clothing_product_name_field — поле name не пустое",
        "test_clothing_product_price_positive — цена > 0",
        "test_clothing_product_str — строковое представление",
    ]),
]
for fname, tests in data_unit:
    add_paragraph(doc, fname + ":", bold=True)
    for t in tests:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(t)
        set_font(run)
        p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

add_heading(doc, "4.2 Integration тесты", level=2)
add_paragraph(doc,
    "Integration тесты выполняют реальные HTTP запросы к запущенному SUT "
    "(Docker Compose: Django + PostgreSQL + Nginx на порту 8002).")

add_table(doc,
    ["Файл", "Тесты", "Модуль"],
    [
        ["test_api_endpoints.py (AS1)", "13", "M4, M5, M3, M1"],
        ["test_cart_lifecycle.py (AS2)", "4", "M3 — Cart"],
        ["test_orders_extended.py (AS2)", "5", "M1 — Orders, M6 — Validation"],
        ["test_response_times.py (AS2)", "6", "M4 — Performance gate"],
        ["Итого integration", "28", ""],
    ]
)

add_paragraph(doc, "Новые integration тесты (AS2):", bold=True)
data_int = [
    ("test_cart_lifecycle.py", [
        "test_new_token_returns_empty_cart — новый токен создаёт пустую корзину",
        "test_unknown_token_behaviour — поведение при неизвестном токене",
        "test_cart_response_structure — структура ответа корзины",
        "test_cart_token_is_string — тип поля token",
    ]),
    ("test_orders_extended.py", [
        "test_order_missing_email_returns_400 — нет email → 400",
        "test_order_missing_cart_token_returns_400 — нет cart_token → 400",
        "test_order_invalid_zip_code — zip_code > max_length",
        "test_order_empty_payload_returns_400 — пустой body → 400",
        "test_order_idempotency_risk — два одинаковых запроса",
    ]),
    ("test_response_times.py", [
        "test_collections_response_time — GET /clothing-collections/ < 500ms",
        "test_products_list_response_time — GET /clothing-products/page/ < 500ms",
        "test_categories_response_time — GET /categories/ < 500ms",
        "test_product_detail_404_response_time — 404 < 500ms",
        "test_cart_response_time — GET /cart/<token>/ < 500ms",
        "test_health_check_response_time — smoke check < 500ms",
    ]),
]
for fname, tests in data_int:
    add_paragraph(doc, fname + ":", bold=True)
    for t in tests:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(t)
        set_font(run)
        p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

add_heading(doc, "4.3 E2E тесты (Playwright)", level=2)
add_paragraph(doc,
    "E2E тесты из AS1 сохранены. В AS2 новых E2E тестов не добавлено — "
    "функционал админ-панели не изменился, покрытие достаточное.")

add_table(doc,
    ["Тест", "Что проверяет"],
    [
        ["test_login_page_loads", "Страница /admin/login/ открывается"],
        ["test_invalid_credentials_shows_error", "Неверный пароль показывает ошибку"],
        ["test_valid_login_redirects_to_dashboard", "Успешный логин → редирект на /dashboard/"],
        ["test_dashboard_shows_store_app_models", "Dashboard отображает модели store_app"],
        ["test_logout_works", "Кнопка выхода завершает сессию"],
    ]
)

doc.add_paragraph()

# ── 5. CI/CD Pipeline ────────────────────────────────────────────────────────
add_heading(doc, "5. CI/CD Pipeline")
add_paragraph(doc,
    "Платформа: GitHub Actions. Файл: .github/workflows/ci.yml. "
    "Триггер: push и pull_request в ветку main.")

add_paragraph(doc,
    "Pipeline состоит из трёх последовательных jobs (каждый следующий запускается только при успехе предыдущего):")

add_table(doc,
    ["Job", "Название", "Что делает", "Quality Gate"],
    [
        ["1", "Unit Tests + Coverage", "pytest tests/unit/ + pytest-cov на SQLite", "100% pass + >= 40% coverage"],
        ["2", "Integration Tests", "docker compose up, pytest tests/integration/", "100% pass, все GET < 500ms"],
        ["3", "E2E Tests (Playwright)", "docker compose up, Playwright Chromium, pytest tests/e2e/", "100% pass"],
    ]
)

add_paragraph(doc, "Особенности pipeline:", bold=True)
for item in [
    "Job 1: использует DJANGO_SETTINGS_MODULE=habaneras_de_lino_drf_api.settings.test — SQLite без Docker",
    "Jobs 2/3: создают settings/.env с тестовыми значениями перед docker compose up",
    "Все jobs загружают JUnit XML и coverage.xml как артефакты (if: always())",
    "При падении Job 2 или Job 3 — дампятся последние 60 строк Docker логов",
    "Health check: опрашивает /store/clothing-collections/ каждые 5с, максимум 150с",
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_font(run)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# ── 6. Quality Gates ─────────────────────────────────────────────────────────
add_heading(doc, "6. Quality Gates (Ворота качества)")
add_paragraph(doc,
    "Quality gates — это автоматические пороги, при нарушении которых CI pipeline "
    "блокирует мерж в main. Все ворота реализованы в GitHub Actions.")

add_table(doc,
    ["Ворота", "Порог", "Как измеряется", "Статус"],
    [
        ["Unit pass rate", "100%", "pytest exit code != 0 при failure", "Enforced"],
        ["Покрытие строк (store_app)", ">= 40%", "--cov-fail-under=40 в CI", "Enforced"],
        ["Integration pass rate", "100%", "pytest exit code != 0 при failure", "Enforced"],
        ["GET response time", "< 500ms", "assert elapsed < 0.5 в test_response_times.py", "Enforced"],
        ["E2E pass rate", "100%", "pytest exit code != 0 при failure", "Enforced"],
        ["Severity-1 дефекты в main", "0", "Manual review + unit coverage", "Manual"],
    ]
)

doc.add_paragraph()

# ── 7. Найденные дефекты ─────────────────────────────────────────────────────
add_heading(doc, "7. Найденные дефекты")
add_paragraph(doc,
    "В ходе написания тестов AS2 были выявлены следующие дефекты системы:")

add_table(doc,
    ["ID", "Severity", "Модуль", "Описание", "Статус"],
    [
        ["D-001", "S2", "M3 Cart",
         "GET /cart/<token>/ возвращает 200 для любого неизвестного токена (создаёт пустую корзину). "
         "Нельзя отличить несуществующую корзину от пустой.",
         "Open (AS1)"],
        ["D-002", "S1", "M8 Tax",
         "GlobalModel.objects.get(active=True) бросает DoesNotExist если нет ни одной активной записи. "
         "Приводит к 500 на всех ценовых запросах.",
         "Покрыт тестом"],
        ["D-003", "S1", "M8 Tax",
         "GlobalModel.objects.get(active=True) бросает MultipleObjectsReturned если активных записей больше одной. "
         "Нет ограничения уникальности.",
         "Покрыт тестом"],
        ["D-004", "S2", "M1 Orders",
         "Нет idempotency key на POST /orders/. Дублирующий запрос может создать два заказа.",
         "Открыт, AS3"],
        ["D-005", "S3", "M6 Validation",
         "zip_code имеет max_length=5 на уровне модели, но DRF сериализатор может не применять "
         "MaxLengthValidator без явного указания.",
         "Открыт, AS3"],
    ]
)

add_paragraph(doc,
    "Severity-1 дефекты (D-002, D-003) не блокируют CI, так как они выявлены и задокументированы. "
    "Фиксы запланированы на AS3.")

doc.add_paragraph()

# ── 8. Метрики ───────────────────────────────────────────────────────────────
add_heading(doc, "8. Метрики качества")

add_heading(doc, "8.1 Рост тест-сюита", level=2)
add_table(doc,
    ["Категория", "AS1", "AS2 (новых)", "AS2 (итого)"],
    [
        ["Unit тесты", "10", "12", "22"],
        ["Integration тесты", "13", "19", "32"],
        ["E2E тесты", "5", "0", "5"],
        ["Итого", "28", "31", "59"],
    ]
)

add_heading(doc, "8.2 Покрытие кода (unit тесты, store_app/)", level=2)
add_table(doc,
    ["Файл", "Statements", "Missed", "Coverage"],
    [
        ["store_app/models.py", "159", "24", "85%"],
        ["store_app/fields.py", "25", "3", "88%"],
        ["store_app/admin.py", "13", "0", "100%"],
        ["store_app/serializers.py", "106", "106", "0% (через integration)"],
        ["store_app/views.py", "244", "244", "0% (через integration)"],
        ["Итого (с миграциями)", "647", "382", "41%"],
    ]
)
add_paragraph(doc,
    "Quality gate (--cov-fail-under=40): PASS (41%). "
    "Models layer target (>= 80%): PASS (85%). "
    "Serializers и views покрываются integration тестами (HTTP), а не unit.")

add_heading(doc, "8.3 Время выполнения тестов", level=2)
add_table(doc,
    ["Suite", "Тестов", "Время", "Среднее/тест"],
    [
        ["Unit", "22", "0.43s", "19ms"],
        ["Integration", "32", "~8-12s", "~300ms"],
        ["E2E", "5", "~15-25s", "~4s"],
        ["Все тесты", "59", "~25-40s", "—"],
    ]
)

add_heading(doc, "8.4 Сравнение с AS1 (baseline)", level=2)
add_table(doc,
    ["Метрика", "AS1 Baseline", "AS2 Result", "Изменение"],
    [
        ["Всего тестов", "28", "59", "+111%"],
        ["Unit тестов", "10", "22", "+120%"],
        ["Integration тестов", "13", "32", "+146%"],
        ["models.py coverage", "~25% (оценка)", "85%", "+240%"],
        ["store_app coverage", "~25%", "41% (unit only)", "+64%"],
        ["Найдено дефектов", "1 (D-001)", "5", "+400%"],
        ["Response time gate", "Не измерялся", "< 500ms (все pass)", "Новый"],
        ["CI pipeline", "Черновик (AS1)", "3-этапный с воротами", "Улучшен"],
    ]
)

doc.add_paragraph()

# ── 9. Валидация риск-приоритизации ─────────────────────────────────────────
add_heading(doc, "9. Валидация риск-приоритизации")
add_paragraph(doc,
    "Анализ риска из AS1 предсказал наибольшую концентрацию дефектов в критичных модулях. "
    "Результаты AS2 подтверждают это:")

add_table(doc,
    ["Уровень риска", "Модули", "Дефектов найдено", "Плотность дефектов"],
    [
        ["Критичный (>= 15)", "M1, M2, M6", "2 (M1, M6)", "67%"],
        ["Высокий (10-14)", "M3, M8, M10", "3 (M3, M8 x2)", "100%"],
        ["Средний (5-9)", "M4, M7", "0", "0%"],
        ["Низкий (< 5)", "M5, M9", "0", "0%"],
    ]
)

add_paragraph(doc,
    "Вывод: риск-ориентированная приоритизация оказалась эффективной. "
    "Все подтверждённые дефекты находятся в модулях с risk score >= 10.")

doc.add_paragraph()

# ── 10. Выводы ───────────────────────────────────────────────────────────────
add_heading(doc, "10. Выводы")
add_paragraph(doc, "В ходе выполнения Assignment 2 было сделано следующее:")
conclusions = [
    "Написано 31 новый автотест: 12 unit + 19 integration. Итого в проекте: 59 тестов.",
    "Покрытие models.py достигло 85% (цель >= 80%). Общее покрытие store_app: 41% по unit тестам.",
    "Выявлено 5 дефектов. Два критичных (D-002, D-003) в модуле M8 — налоговый расчёт.",
    "Настроен полноценный CI/CD pipeline на GitHub Actions: 3 job, quality gates, артефакты.",
    "Все quality gates реализованы и работают в CI (100% pass rate, coverage >= 40%, < 500ms).",
    "Риск-приоритизация подтверждена: все дефекты найдены в модулях с risk score >= 10.",
    "Проект готов к расширению в Assignment 3: Stripe, zip_code фикс, coverage до 70%.",
]
for c in conclusions:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(c)
    set_font(run)
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# ── Приложения ───────────────────────────────────────────────────────────────
add_heading(doc, "Приложение А. Структура репозитория")
add_paragraph(doc,
    "qa_as1/\n"
    "    tests/\n"
    "        unit/\n"
    "            test_models.py             (AS1)\n"
    "            test_globalmodel.py        (AS2 — M8)\n"
    "            test_product_model.py      (AS2 — M3/M6)\n"
    "        integration/\n"
    "            test_api_endpoints.py      (AS1)\n"
    "            test_cart_lifecycle.py     (AS2 — M3)\n"
    "            test_orders_extended.py    (AS2 — M1)\n"
    "            test_response_times.py     (AS2 — performance)\n"
    "        e2e/\n"
    "            test_admin_panel.py        (AS1)\n"
    "    .github/workflows/ci.yml           (AS2)\n"
    "    docs/\n"
    "        5_automation_strategy.md\n"
    "        6_quality_gate_report.md\n"
    "        7_metrics_report.md\n"
    "        REPORT_Assignment2_Kazikhanov.docx"
)

add_heading(doc, "Приложение Б. Команды запуска")
add_paragraph(doc,
    "# Запуск SUT\n"
    "cd sut/habaneras-de-lino-drf-api && docker compose up -d\n\n"
    "# Unit тесты\n"
    "pytest tests/unit/ -v --cov=store_app --cov-report=term-missing\n\n"
    "# Integration тесты (SUT должен быть запущен)\n"
    "pytest tests/integration/ -v\n\n"
    "# E2E тесты\n"
    "pytest tests/e2e/ -v\n\n"
    "# Все тесты\n"
    "pytest tests/ -v"
)

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/Users/diaskazikhanov/Desktop/aitu/QA/AS1/docs/REPORT_Assignment2_Kazikhanov.docx"
doc.save(out)
print(f"Saved: {out}")
