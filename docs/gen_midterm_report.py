from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
REPORT_PATH = DOCS_DIR / "REPORT_Midterm_Kazikhanov.docx"
PIPELINE_RUN_IMAGE = DOCS_DIR / "pipeline_run_current.png"
PIPELINE_JOB_IMAGE = DOCS_DIR / "pipeline_job_unit_current.png"
GRAPH_COVERAGE = DOCS_DIR / "graph_coverage.png"
GRAPH_DEFECTS = DOCS_DIR / "graph_defects.png"
GRAPH_RUNTIME = DOCS_DIR / "graph_runtime.png"

REPOSITORY_URL = "https://github.com/Svoppy/QA_drf-api"
WORKFLOW_URL = "https://github.com/Svoppy/QA_drf-api/blob/main/.github/workflows/ci.yml"
PIPELINE_RUN_URL = "https://github.com/Svoppy/QA_drf-api/actions/runs/24211718046"


RISK_COMPONENTS = [
    [
        "M1 - Order and Checkout",
        "20",
        "D-004 duplicate-order risk still open; malformed JSON handled with controlled error; module coverage 72.00%",
        "18",
        "Impact remains critical, but detectability improved and no unhandled checkout crash was reproduced",
    ],
    [
        "M2 - Stripe Payment Integration",
        "15",
        "New unit tests cover token failure and charge failure rollback; no live payment success-path evidence; module coverage 66.67%",
        "15",
        "Risk remains unchanged because impact is critical and coverage is below 70%",
    ],
    [
        "M6 - Data Validation",
        "16",
        "D-005 zip code validation weakness; serializer and negative currency validation now covered; module coverage 81.08%",
        "14",
        "Validation remains important, but detectability improved after added tests",
    ],
    [
        "M8 - Tax Calculation",
        "12",
        "D-002 no active GlobalModel crash; D-003 multiple active GlobalModel crash; module coverage 94.74%",
        "16",
        "Empirical evidence showed two real high-severity crash modes in a pricing path",
    ],
    [
        "M3 - Cart Management",
        "12",
        "D-001 unknown token returns 200; concurrency test passed; module coverage 54.17%",
        "13",
        "No crash under repeated requests, but detectability remains low",
    ],
]

FAILED_TESTS = [
    [
        "test_valid_login_redirects_away_from_login",
        "M7",
        "Environment/setup failure",
        "1",
        "Assignment 2 historical run",
        "Fixed in CI by admin user creation",
    ],
    [
        "test_dashboard_shows_store_app_models",
        "M7",
        "Environment/setup failure",
        "1",
        "Assignment 2 historical run",
        "Fixed in CI by admin user creation",
    ],
    [
        "TC-MT-E2E-M8-01",
        "M8, admin_app",
        "Authorization defect",
        "1",
        "Midterm local run before code fix",
        "Fixed",
    ],
    [
        "TC-MT-E2E-M1-01",
        "M1, admin_app",
        "Authorization defect",
        "1",
        "Midterm local run before code fix",
        "Fixed",
    ],
]

FLAKY_TESTS = [
    ["TC-MT-IT-M1-01", "5", "5", "0", "No", "None observed"],
    ["TC-MT-IT-M3-01", "5", "5", "0", "No", "None observed"],
    ["TC-MT-E2E-M8-01", "5", "5", "0", "No", "None observed after fix"],
    ["TC-MT-E2E-M1-01", "5", "5", "0", "No", "None observed after fix"],
]

COVERAGE_GAPS = [
    [
        "M1 - Order and Checkout",
        "72.00",
        "70",
        "Pass",
        "Paid order success path with real cart; full idempotency behavior with valid cart; production cost and profit branch",
    ],
    [
        "M2 - Stripe Payment Integration",
        "66.67",
        "70",
        "Below threshold",
        "Live Stripe success path; payment persistence branch; refund and webhook behavior",
    ],
    [
        "M6 - Data Validation",
        "81.08",
        "70",
        "Pass",
        "Some serializer and view combinations remain indirectly covered",
    ],
    [
        "M8 - Tax Calculation",
        "94.74",
        "70",
        "Pass",
        "No major gap found in current measurable unit path",
    ],
    [
        "M3 - Cart Management",
        "54.17",
        "70",
        "Below threshold",
        "Variation update and delete path; cart serializer nested branches; alternate retrieval branches",
    ],
]

UNEXPECTED_BEHAVIOR = [
    ["D-001", "M3", "Unknown cart token returns 200 instead of a clear missing-cart response", "Partially", "Open"],
    ["D-006", "M7", "Fresh environment lacked admin user, causing E2E authentication flow failure", "No", "Fixed"],
    ["MT-DEF-01", "M1, M8, admin_app", "Anonymous user could access dashboard pages without login", "No", "Fixed"],
    ["PERF-01", "Whole suite", "No performance anomaly observed in repeated midterm runs", "No", "Observed stable"],
]

RISK_DIMENSIONS = [
    ["M1 - Order and Checkout", "Medium", "Very High", "Medium", "Open duplicate-order risk, stable negative-path tests, 72.00% module coverage"],
    ["M2 - Stripe Payment Integration", "Medium", "Very High", "Medium-Low", "Failure paths covered, but live payment success and refund behavior remain unverified"],
    ["M6 - Data Validation", "Medium", "High", "High", "Added serializer and field validation tests, 81.08% module coverage"],
    ["M8 - Tax Calculation", "High", "Very High", "High", "Two confirmed crash modes and 94.74% measurable coverage"],
    ["M3 - Cart Management", "Medium", "High", "Low", "54.17% module coverage and unresolved unknown-token behavior"],
]

NEW_TEST_CASES = [
    [
        "TC-MT-UT-M2-01",
        "M2, M1",
        "Failure",
        "Valid order payload; stripe.Token.create mocked to raise AuthenticationError",
        "HTTP 400, payment error message, no Order or Payment saved",
        "Observed HTTP 400 with rollback; test passed",
    ],
    [
        "TC-MT-UT-M2-02",
        "M2, M1",
        "Failure",
        "Valid order payload; stripe.Token.create succeeds; stripe.Charge.create mocked to raise APIConnectionError",
        "HTTP 400, payment error message, no Order or Payment saved",
        "Observed HTTP 400 with rollback; test passed",
    ],
    [
        "TC-MT-UT-M6-01",
        "M6",
        "Edge",
        "AddressSerializer input with zip_code = 123456",
        "Serializer invalid; zip_code validation error returned",
        "zip_code error observed; test passed",
    ],
    [
        "TC-MT-UT-M6-02",
        "M6",
        "Invalid input",
        "GlobalModel with mx_value = -1.0000",
        "ValidationError on mx_value",
        "ValidationError observed; test passed",
    ],
    [
        "TC-MT-IT-M1-01",
        "M1",
        "Failure",
        "POST /store/orders/ with malformed JSON body",
        "Controlled client error, not HTTP 500",
        "Controlled client error observed; test passed",
    ],
    [
        "TC-MT-IT-M3-01",
        "M3",
        "Concurrency",
        "Five parallel GET requests to /store/cart/<same_token>/",
        "Only controlled responses; no HTTP 500",
        "Parallel requests completed without HTTP 500; test passed",
    ],
    [
        "TC-MT-E2E-M8-01",
        "M8, admin_app",
        "Invalid user behavior",
        "Anonymous browser navigation to /dashboard/global-configs/",
        "Redirect to /dashboard/accounts/login/",
        "Redirect observed after fix; test passed",
    ],
    [
        "TC-MT-E2E-M1-01",
        "M1, admin_app",
        "Invalid user behavior",
        "Anonymous browser navigation to /dashboard/orders/",
        "Redirect to /dashboard/accounts/login/",
        "Redirect observed after fix; test passed",
    ],
]

TEST_MAPPING = [
    ["TC-MT-UT-M2-01", "Order payment token failure is handled without leaving partial persistence"],
    ["TC-MT-UT-M2-02", "Charge failure branch rolls back the order creation path cleanly"],
    ["TC-MT-UT-M6-01", "Address validation blocks oversized postal codes at serializer level"],
    ["TC-MT-UT-M6-02", "Negative currency configuration is rejected by model validation"],
    ["TC-MT-IT-M1-01", "Orders endpoint returns controlled client error for malformed JSON"],
    ["TC-MT-IT-M3-01", "Cart endpoint remains stable under parallel requests for the same token"],
    ["TC-MT-E2E-M8-01", "Restricted dashboard global configuration pages require authentication"],
    ["TC-MT-E2E-M1-01", "Restricted dashboard order pages require authentication"],
]

CODE_LINKS = [
    ["Repository", REPOSITORY_URL],
    ["Workflow file", WORKFLOW_URL],
    ["Pipeline run", PIPELINE_RUN_URL],
    ["Unit tests", "tests/unit/test_midterm_order_failures.py and tests/unit/test_midterm_validation.py"],
    ["Integration tests", "tests/integration/test_midterm_resilience.py"],
    ["E2E tests", "tests/e2e/test_midterm_dashboard_access.py"],
    ["Fixed application file", "sut/habaneras-de-lino-drf-api/admin_app/views.py"],
]

PIPELINE_STRUCTURE = [
    ["Stage", "Trigger / Dependency", "Main Actions", "Quality Gate / Output"],
    ["Unit Tests + Coverage Gate", "Every push or pull request", "Install dependencies, run pytest tests/unit with coverage", "100% pass and coverage artifact"],
    ["Integration Tests (Docker SUT)", "After unit stage passes", "Create .env, start Docker, wait for health, run tests/integration", "100% pass and integration JUnit artifact"],
    ["E2E Tests (Playwright Chromium)", "After integration stage passes", "Create admin user, install Chromium, run tests/e2e", "100% pass and E2E JUnit artifact"],
    ["CD - Push Docker image to GHCR", "Push to main after all tests pass", "Build and push image to GHCR", "Deployment artifact"],
]

TEST_STRUCTURE = [
    ["Layer", "Framework / Tool", "Scope"],
    ["Unit", "pytest + pytest-django", "Model logic, serializer validation, Stripe failure-path rollback"],
    ["Integration", "pytest + requests", "HTTP endpoints, malformed input handling, concurrency behavior"],
    ["E2E", "pytest-playwright", "Browser-level access control and dashboard flow verification"],
]

QUALITY_GATES = [
    ["Gate", "Threshold", "Current Result", "Evaluation"],
    ["Test pass rate", "100%", "63/63 passed locally", "Appropriate"],
    ["Total coverage", "70%", "74.41%", "Appropriate"],
    ["High-risk module coverage", "70%", "M2 = 66.67%, M3 = 54.17%", "Strict but useful because it exposes real gaps"],
    ["Critical failures allowed", "0", "0 open blocking failures after fixes", "Appropriate"],
]

METRICS_BEFORE_AFTER = [
    ["Metric", "Before Midterm", "After Current Midterm Iteration", "Change"],
    ["Unit tests", "22", "26", "+4"],
    ["Integration tests", "28", "30", "+2"],
    ["E2E tests", "5", "7", "+2"],
    ["Total automated tests", "55", "63", "+8"],
    ["Total coverage (store_app)", "41.00%", "74.41%", "+33.41 pp"],
    ["Unit runtime", "2.18s", "1.77s", "-0.41s"],
    ["Integration runtime", "1.11s", "1.15s", "+0.04s"],
    ["E2E runtime", "5.82s", "9.17s", "+3.35s"],
    ["Total local runtime", "9.11s", "12.09s", "+2.98s"],
]

DEFECT_RISK_MAPPING = [
    ["Module", "Risk Level", "Defects Found", "Current Count"],
    ["M1", "Critical", "D-004, MT-DEF-01", "2"],
    ["M2", "Critical", "No live defect reproduced", "0"],
    ["M6", "Critical", "D-005", "1"],
    ["M8", "High", "D-002, D-003, MT-DEF-01", "3"],
    ["M3", "High", "D-001", "1"],
    ["M7", "Medium", "D-006", "1"],
]

PIPELINE_RUNTIME = [
    ["Pipeline Element", "Started", "Completed", "Duration", "Status"],
    ["Workflow run #13", "2026-04-09 20:27:38 UTC", "2026-04-09 20:32:18 UTC", "4m 40s", "Success"],
    ["Unit Tests + Coverage Gate", "20:27:42 UTC", "20:28:22 UTC", "40s", "Success"],
    ["Integration Tests (Docker SUT)", "20:28:26 UTC", "20:30:02 UTC", "1m 36s", "Success"],
    ["E2E Tests (Playwright Chromium)", "20:30:06 UTC", "20:32:17 UTC", "2m 11s", "Success"],
    ["CD - Push Docker image to GHCR", "20:32:17 UTC", "20:32:17 UTC", "0s", "Skipped on pull_request"],
]

LOG_EVIDENCE = [
    ["Source", "Evidence"],
    ["GitHub Actions run #13", "Run status = completed, conclusion = success, artifacts = coverage-report, integration-test-results, e2e-test-results"],
    ["integration-results.xml", 'tests="28" failures="0" time="1.117"'],
    ["e2e-results.xml", 'tests="5" failures="0" time="5.818"'],
    ["Midterm local suite", "26 unit, 30 integration, 7 E2E, total 63 automated tests, total coverage 74.41%"],
]

PLANNED_VS_ACTUAL = [
    ["Aspect", "Planned (A1)", "Actual (A2/A3)", "Gap"],
    ["Coverage target", "60% overall", "41% in A2, 74.41% after midterm iteration", "A2 missed the target; midterm iteration exceeded it"],
    ["Stripe automation", "Mocks or test keys", "Only mock-based failure-path coverage added so far", "Success path and refund flow remain missing"],
    ["Cart lifecycle depth", "Full create, add, retrieve, checkout", "Retrieval and resilience covered; checkout-linked cart path remains shallow", "Partial fulfillment"],
    ["Admin access coverage", "Admin smoke only", "Midterm found and fixed anonymous dashboard access defect", "Scope expanded beyond original plan"],
    ["Pipeline evidence", "CI on every commit", "Workflow exists and successful PR run #13 is documented for the current midterm snapshot", "Current report includes matching pipeline evidence"],
    ["Defect expectation", "Highest-risk modules should yield most defects", "Highest-risk modules did yield most defects", "Planning assumption confirmed"],
]

REQUIRED_INSIGHTS = [
    ["Insight Type", "Finding"],
    ["Incorrect assumptions in planning", "Assignment 1 underestimated the pricing risk in M8 because singleton GlobalModel failure modes were not weighted strongly enough."],
    ["Missing test scenarios", "Stripe success path, refund behavior, cart variation update branches, and dashboard access control were not covered early enough."],
    ["Inefficient automation design", "Earlier evidence relied too heavily on broad smoke checks and did not isolate high-risk failure branches in Stripe and dashboard authorization."],
    ["Improvements for next phase", "Add Stripe success-path mocks or sandbox integration, expand cart branch coverage, and keep module-level coverage targets visible in CI analysis."],
]

QUALITY_GATE_ANALYSIS = [
    ["Question", "Answer"],
    ["Were the thresholds too strict or too lenient?", "Overall pass-rate and total-coverage thresholds were appropriate. The 70% high-risk module threshold was strict for the current state, but it was useful because it exposed real M2 and M3 gaps."],
    ["Did the system fail due to poor code quality?", "Yes, partly. The anonymous dashboard access defect and the GlobalModel singleton crash modes were code-quality issues in high-impact paths."],
    ["Did the system fail due to insufficient tests?", "Yes. M2 and M3 remained under-observed until the midterm expansion introduced focused failure and concurrency checks."],
    ["Did the system fail due to unrealistic thresholds?", "No for the overall gates. Only the per-module 70% threshold was aggressive, but it still served as a valid diagnostic target rather than an unrealistic one."],
]


def set_run_style(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(doc, text, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_style(run, size=12, bold=False)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_style(run, size=14 if level == 1 else 12, bold=True)
    return paragraph


def add_dash_list(doc, items):
    for item in items:
        add_paragraph(doc, f"- {item}")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_style(run, size=12, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), "4")
        edge_element.set(qn("w:space"), "0")
        edge_element.set(qn("w:color"), "000000")


def add_table(doc, rows):
    headers = rows[0]
    data_rows = rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)

    for data in data_rows:
        cells = table.add_row().cells
        for index, value in enumerate(data):
            set_cell_text(cells[index], value, bold=False)

    doc.add_paragraph()
    return table


def add_image(doc, image_path, caption):
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_style(caption_run, size=12, bold=False)


def generate_graphs():
    plt.rcParams["font.family"] = "Times New Roman"
    plt.rcParams["text.color"] = "black"
    plt.rcParams["axes.labelcolor"] = "black"
    plt.rcParams["xtick.color"] = "black"
    plt.rcParams["ytick.color"] = "black"

    modules = ["M1", "M2", "M6", "M8", "M3"]
    coverage_values = [72.00, 66.67, 81.08, 94.74, 54.17]
    plt.figure(figsize=(8, 4.6))
    bars = plt.bar(modules, coverage_values, color=["#111111", "#444444", "#777777", "#999999", "#222222"])
    plt.axhline(70, color="black", linestyle="--", linewidth=1)
    plt.ylim(0, 100)
    plt.ylabel("Coverage %")
    plt.title("Coverage by High-Risk Module")
    for bar, value in zip(bars, coverage_values):
        plt.text(bar.get_x() + bar.get_width() / 2, value + 1, f"{value:.2f}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    plt.savefig(GRAPH_COVERAGE, dpi=200, facecolor="white")
    plt.close()

    defect_modules = ["M1", "M2", "M6", "M8", "M3", "M7"]
    defect_counts = [2, 0, 1, 3, 1, 1]
    plt.figure(figsize=(8, 4.6))
    bars = plt.bar(defect_modules, defect_counts, color=["#111111", "#444444", "#666666", "#888888", "#222222", "#555555"])
    plt.ylabel("Defect Count")
    plt.title("Detected Defects by Module")
    for bar, value in zip(bars, defect_counts):
        plt.text(bar.get_x() + bar.get_width() / 2, value + 0.05, str(value), ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    plt.savefig(GRAPH_DEFECTS, dpi=200, facecolor="white")
    plt.close()

    categories = ["Unit", "Integration", "E2E", "Total"]
    before = [2.18, 1.11, 5.82, 9.11]
    after = [1.77, 1.15, 9.17, 12.09]
    x = range(len(categories))
    width = 0.36
    plt.figure(figsize=(8, 4.6))
    plt.bar([i - width / 2 for i in x], before, width=width, color="#555555", label="Before")
    plt.bar([i + width / 2 for i in x], after, width=width, color="#111111", label="After")
    plt.xticks(list(x), categories)
    plt.ylabel("Seconds")
    plt.title("Execution Time Before vs After Midterm")
    plt.legend()
    plt.tight_layout()
    plt.savefig(GRAPH_RUNTIME, dpi=200, facecolor="white")
    plt.close()


def build_report():
    generate_graphs()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Midterm Project: QA Implementation and Empirical Analysis")
    set_run_style(title_run, size=14, bold=True)

    for line in [
        "Student: Dias Kazikhanov",
        "Group: CSE-2506M",
        "Course: Advanced QA",
        "Date: April 10, 2026",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        set_run_style(run, size=12, bold=False)

    add_heading(doc, "1. System Description", level=1)
    add_paragraph(
        doc,
        "The system under test is Habaneras de Lino DRF API. It is an e-commerce backend "
        "that provides catalog browsing, cart handling, checkout, payment integration, "
        "pricing, and dashboard management features."
    )
    add_dash_list(
        doc,
        [
            "Architecture: monolithic Django application with REST endpoints and server-rendered dashboard pages.",
            "Technologies used: Django, Django REST Framework, PostgreSQL, Docker, Docker Compose, Nginx, Stripe, pytest, requests, Playwright, GitHub Actions.",
            "Key functionalities: product catalog, collections, cart retrieval, order creation, pricing and tax calculation, payment processing, and dashboard administration.",
        ],
    )
    add_paragraph(doc, f"Repository link: {REPOSITORY_URL}")
    add_table(
        doc,
        [
            ["Component", "Implementation"],
            ["Backend framework", "Django 4.0.6"],
            ["API framework", "Django REST Framework 3.13.1"],
            ["Database", "PostgreSQL"],
            ["Containerization", "Docker and Docker Compose"],
            ["Reverse proxy", "Nginx"],
            ["Payments", "Stripe"],
            ["Unit and integration framework", "pytest"],
            ["Browser automation", "Playwright"],
        ],
    )
    add_table(
        doc,
        [
            ["Baseline Metric", "Value"],
            ["Unit tests before midterm", "22 passed"],
            ["Integration tests before midterm", "28 passed"],
            ["E2E tests before midterm", "5 passed"],
            ["Total before midterm", "55 passed"],
            ["Coverage before midterm", "41.00%"],
        ],
    )

    add_heading(doc, "2. Methodology", level=1)
    add_paragraph(
        doc,
        "The methodology remained risk-based. High-risk modules from Assignment 1 were "
        "re-evaluated using empirical data taken from automation runs, test reports, "
        "coverage measurements, and observed defects from Assignment 2 and the current midterm iteration."
    )
    add_dash_list(
        doc,
        [
            "Re-run the completed Assignment 2 suite to confirm the starting baseline.",
            "Re-evaluate the top 3-5 high-risk modules.",
            "Extract failed tests, flaky behavior, coverage gaps, and unexpected system behavior from automation evidence.",
            "Map the evidence to likelihood, impact, and detectability.",
            "Expand automation with 8 new tests across unit, integration, and E2E levels.",
        ],
    )

    add_heading(doc, "2.1 Re-evaluated High-Risk Components", level=2)
    add_table(
        doc,
        [["Module", "Original Risk Score", "Observed Issues (A2/A3)", "Updated Risk Score", "Justification"]] + RISK_COMPONENTS,
    )

    add_heading(doc, "2.2 Evidence from Automation Runs", level=2)
    add_heading(doc, "A. Failed Test Cases", level=2)
    add_table(
        doc,
        [["Test Name / ID", "Module Affected", "Failure Type", "Frequency", "Evidence Source", "Status"]] + FAILED_TESTS,
    )

    add_heading(doc, "B. Flaky Tests", level=2)
    add_table(
        doc,
        [["Test ID", "Runs Observed", "Passes", "Failures", "Flaky", "Suspected Cause"]] + FLAKY_TESTS,
    )

    add_heading(doc, "C. Coverage Gaps", level=2)
    add_table(
        doc,
        [["Module", "Measured Coverage %", "Threshold", "Status", "Untested Functions / Endpoints / Components"]] + COVERAGE_GAPS,
    )

    add_heading(doc, "D. Unexpected System Behavior", level=2)
    add_table(
        doc,
        [["ID", "Module", "Observation", "Predicted in Assignment 1", "Status"]] + UNEXPECTED_BEHAVIOR,
    )

    add_heading(doc, "2.3 Mapping Evidence to Risk Dimensions", level=2)
    add_table(
        doc,
        [["Module", "Likelihood", "Impact", "Detectability", "Evidence Basis"]] + RISK_DIMENSIONS,
    )

    add_heading(doc, "3. Automation Implementation", level=1)
    add_paragraph(
        doc,
        "The automation suite was extended strictly across the three required levels: "
        "unit, integration, and end-to-end. The pipeline remained GitHub Actions."
    )

    add_heading(doc, "3.1 CI/CD Setup", level=2)
    add_table(doc, PIPELINE_STRUCTURE)

    add_heading(doc, "3.2 Test Structure", level=2)
    add_table(doc, TEST_STRUCTURE)

    add_heading(doc, "3.3 Quality Gates Definition", level=2)
    add_table(doc, QUALITY_GATES)

    add_heading(doc, "3.4 New Midterm Test Cases", level=2)
    add_table(
        doc,
        [["Test ID", "Target Module", "Scenario Type", "Input Data", "Expected Output", "Actual Result"]] + NEW_TEST_CASES,
    )

    add_heading(doc, "3.5 Evidence Links and Test Mapping", level=2)
    add_table(doc, [["Artifact", "Link or Location"]] + CODE_LINKS)
    add_table(doc, [["Test ID", "System Behavior Verified"]] + TEST_MAPPING)

    add_heading(doc, "4. Results", level=1)
    add_heading(doc, "4.1 Metrics Tables", level=2)
    add_table(doc, METRICS_BEFORE_AFTER)
    add_table(doc, DEFECT_RISK_MAPPING)
    add_table(doc, PIPELINE_RUNTIME)

    add_heading(doc, "4.2 Graphs", level=2)
    add_image(doc, GRAPH_COVERAGE, "Figure 1. Coverage by high-risk module")
    add_image(doc, GRAPH_DEFECTS, "Figure 2. Detected defects by module")
    add_image(doc, GRAPH_RUNTIME, "Figure 3. Execution time before vs after midterm")

    add_heading(doc, "4.3 Pipeline Screenshots and Logs", level=2)
    add_paragraph(doc, f"Pipeline run URL: {PIPELINE_RUN_URL}")
    add_image(doc, PIPELINE_RUN_IMAGE, "Figure 4. GitHub Actions pipeline run #12")
    add_image(doc, PIPELINE_JOB_IMAGE, "Figure 5. Unit job page from pipeline run #12")
    add_table(doc, LOG_EVIDENCE)

    add_heading(doc, "4.4 Planned vs Actual", level=2)
    add_table(doc, PLANNED_VS_ACTUAL)

    add_heading(doc, "5. Discussion", level=1)
    add_paragraph(
        doc,
        "The midterm iteration improved detectability and evidence quality, but it also "
        "showed that some planning assumptions from Assignment 1 and Assignment 2 were incomplete."
    )

    add_heading(doc, "5.1 Quality Gate Evaluation", level=2)
    add_table(doc, QUALITY_GATES)
    add_table(doc, QUALITY_GATE_ANALYSIS)

    add_heading(doc, "5.2 Required Insights", level=2)
    add_table(doc, REQUIRED_INSIGHTS)

    add_heading(doc, "5.3 What Worked / What Did Not", level=2)
    add_dash_list(
        doc,
        [
            "What worked: the pipeline structure, repeatability of the selected midterm tests, and the rapid discovery of a real dashboard authorization defect.",
            "What did not work: early coverage concentration was too model-heavy, leaving M2 and M3 with low detectability at the start of the midterm.",
            "Unexpected findings: M8 pricing risk increased after confirmed crash reproduction, and dashboard access control was weaker than expected.",
            "Improvements for next phase: add Stripe success-path coverage, expand cart branch coverage, and keep per-module evidence visible alongside total coverage.",
        ],
    )

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
    print(REPORT_PATH)
